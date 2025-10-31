#!/usr/bin/env python3
"""
ABOUTME: Production-grade academic document export utility with strategy pattern
ABOUTME: Supports multiple PDF engines (LibreOffice, Pandoc, WeasyPrint) with auto-fallback
"""

import sys
import argparse
from pathlib import Path
from typing import Optional, Literal

from pdf_engines import (
    PDFGenerationOptions,
    PDFEngineFactory,
    get_available_engines,
    get_recommended_engine
)


def export_pdf(
    md_file: Path,
    output_pdf: Path,
    engine: Literal['auto', 'libreoffice', 'pandoc', 'weasyprint'] = 'auto',
    options: Optional[PDFGenerationOptions] = None
) -> bool:
    """
    Export markdown to PDF using specified engine.

    Uses strategy pattern with automatic fallback to other engines on failure.
    This follows SOLID principles and provides production-grade error handling.

    Args:
        md_file: Path to input markdown file
        output_pdf: Path for output PDF file
        engine: PDF engine to use ('auto' for automatic selection)
        options: PDF generation options (uses academic defaults if None)

    Returns:
        bool: True if PDF generation succeeded, False otherwise

    Examples:
        >>> export_pdf(Path('thesis.md'), Path('thesis.pdf'))
        >>> export_pdf(Path('thesis.md'), Path('thesis.pdf'), engine='libreoffice')
        >>> options = PDFGenerationOptions(line_spacing=1.5)
        >>> export_pdf(Path('thesis.md'), Path('thesis.pdf'), options=options)
    """
    if options is None:
        options = PDFGenerationOptions()

    print(f"\n📄 Generating PDF: {output_pdf.name}")
    print(f"   Input: {md_file}")
    print(f"   Engine: {engine}")
    print()

    # Use factory to generate with automatic fallback
    result = PDFEngineFactory.generate_with_fallback(
        md_file=md_file,
        output_pdf=output_pdf,
        options=options,
        preferred_engine=engine if engine != 'auto' else None
    )

    # Display result
    if result.success:
        print(f"\n✅ PDF created successfully: {result.output_path}")
        print(f"   Engine used: {result.engine_name}")

        if result.warnings:
            print("\n⚠️  Warnings:")
            for warning in result.warnings:
                print(f"   - {warning}")

        return True
    else:
        print(f"\n❌ PDF generation failed")
        print(f"   Error: {result.error_message}")
        return False


def export_docx_basic(md_file: Path, output_docx: Path) -> bool:
    """
    Export markdown to DOCX format using basic python-docx parsing.

    Note: This is a legacy/fallback function with limited formatting support.
    Does NOT support tables, complex formatting, or citations.

    For production use, prefer export_docx() which uses Pandoc with full
    markdown support and proper table formatting.

    Args:
        md_file: Path to input markdown file
        output_docx: Path for output DOCX file

    Returns:
        bool: True if DOCX generation succeeded, False otherwise
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ Error: python-docx not installed")
        print("   Run: pip install python-docx")
        return False

    try:
        # Read markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # Create document
        doc = Document()

        # Set margins (1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Add footer instruction for page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("[Add page numbers via Insert > Page Number in Word]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.italic = True

        # Process markdown content
        for line in lines:
            line = line.rstrip()

            if not line:
                continue

            # Title (# heading)
            if line.startswith('# '):
                para = doc.add_heading(line[2:], level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Section heading (## heading)
            if line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                continue

            # Subsection heading (### heading)
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                continue

            # Horizontal rule
            if line.startswith('---'):
                para = doc.add_paragraph()
                run = para.add_run('_' * 60)
                run.font.size = Pt(12)
                continue

            # Regular paragraph
            para = doc.add_paragraph(line)
            para_format = para.paragraph_format
            para_format.line_spacing = 2.0  # Double spacing
            para_format.space_after = Pt(0)

            # Set font
            if para.runs:
                run = para.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Save
        doc.save(output_docx)

        print(f"✅ DOCX created successfully: {output_docx}")
        return True

    except Exception as e:
        print(f"❌ DOCX generation failed: {str(e)}")
        return False


def export_docx(
    md_file: Path,
    output_docx: Path,
    options: Optional[PDFGenerationOptions] = None
) -> bool:
    """
    Export markdown to DOCX with full formatting support (tables, citations, styles).

    Uses Pandoc with custom reference document for professional formatting that matches
    PDF output quality. Automatically falls back to basic export if Pandoc unavailable.

    Features:
    - ✅ Tables rendered as Word tables (not pipe text)
    - ✅ Proper heading styles (APA 7th edition compatible)
    - ✅ Citations and footnotes
    - ✅ Table of contents
    - ✅ Custom fonts and spacing from reference document

    Args:
        md_file: Path to input markdown file
        output_docx: Path for output DOCX file
        options: PDF generation options (uses academic defaults if None)

    Returns:
        bool: True if DOCX generation succeeded, False otherwise

    Examples:
        >>> export_docx(Path('thesis.md'), Path('thesis.docx'))
        >>> options = PDFGenerationOptions(enable_toc=True, toc_depth=3)
        >>> export_docx(Path('thesis.md'), Path('thesis.docx'), options)
    """
    import subprocess
    import shutil

    if options is None:
        options = PDFGenerationOptions()

    # Try Pandoc method first (best quality)
    if not shutil.which('pandoc'):
        print("⚠️  Pandoc not found - falling back to basic DOCX export")
        print("   (Tables and advanced formatting will be limited)")
        print("   Install Pandoc for better results: sudo apt install pandoc")
        return export_docx_basic(md_file, output_docx)

    try:
        # Locate reference document
        reference_doc = Path(__file__).parent.parent / "examples" / "custom-reference.docx"

        if not reference_doc.exists():
            print(f"⚠️  Warning: Reference document not found: {reference_doc}")
            print("   Continuing without reference document (may lose some formatting)")
            reference_doc = None

        # Build pandoc command
        cmd = [
            'pandoc',
            str(md_file),
            '-o', str(output_docx),
            '--from', 'markdown',
            '--to', 'docx',
        ]

        # Add reference document for styling
        if reference_doc:
            cmd.extend(['--reference-doc', str(reference_doc)])

        # Add table of contents
        if options.enable_toc:
            cmd.append('--toc')
            cmd.extend(['--toc-depth', str(options.toc_depth)])

        # Add metadata if provided
        if options.title:
            cmd.extend(['--metadata', f'title={options.title}'])
        if options.author:
            cmd.extend(['--metadata', f'author={options.author}'])

        print(f"\n📄 Generating DOCX with Pandoc: {output_docx.name}")
        print(f"   Input: {md_file}")
        if reference_doc:
            print(f"   Reference doc: {reference_doc.name}")
        print()

        # Run pandoc
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            print(f"❌ Pandoc failed with return code {result.returncode}")
            if result.stderr:
                print(f"   Error: {result.stderr}")
            return False

        print(f"✅ DOCX created successfully: {output_docx}")
        print(f"   Tables, formatting, and styling preserved from markdown")
        return True

    except subprocess.TimeoutExpired:
        print("❌ DOCX generation timed out (>60s)")
        return False
    except Exception as e:
        print(f"❌ DOCX generation failed: {str(e)}")
        return False


def show_available_engines() -> None:
    """Display available PDF engines and their status."""
    print("\n📋 Available PDF Engines:")
    print()

    engines = get_available_engines()
    if not engines:
        print("   ❌ No PDF engines available")
        print()
        print("   Install at least one of:")
        print("   - LibreOffice: sudo apt install libreoffice-writer libreoffice-core-nogui")
        print("   - Pandoc: sudo apt install pandoc texlive-latex-base texlive-latex-recommended")
        print("   - WeasyPrint: pip install weasyprint")
        return

    recommended = get_recommended_engine()

    for engine in engines:
        marker = "⭐" if engine == recommended else "  "
        print(f"   {marker} {engine}")

    print()
    if recommended:
        print(f"   Recommended: {recommended}")
    print()


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description='Export markdown to professional academic PDF/DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Auto-select best engine
  python export_professional.py thesis.md --pdf thesis.pdf

  # Use specific engine
  python export_professional.py thesis.md --pdf thesis.pdf --engine libreoffice

  # Generate both PDF and DOCX
  python export_professional.py thesis.md --pdf thesis.pdf --docx thesis.docx

  # List available engines
  python export_professional.py --list-engines

Engine Options:
  auto         - Automatically select best available engine (recommended)
  libreoffice  - Use LibreOffice (best quality, requires installation)
  pandoc       - Use Pandoc/LaTeX (highest quality, academic standard)
  weasyprint   - Use WeasyPrint (fallback, known font rendering issues)
        '''
    )
    parser.add_argument('input', nargs='?', help='Input markdown file')
    parser.add_argument('--pdf', help='Output PDF file')
    parser.add_argument('--docx', help='Output DOCX file')
    parser.add_argument(
        '--engine',
        choices=['auto', 'libreoffice', 'pandoc', 'weasyprint'],
        default='auto',
        help='PDF engine to use (default: auto)'
    )
    parser.add_argument(
        '--list-engines',
        action='store_true',
        help='List available PDF engines and exit'
    )

    args = parser.parse_args()

    # List engines mode
    if args.list_engines:
        show_available_engines()
        sys.exit(0)

    # Validate arguments
    if not args.input:
        parser.error("Input file required (or use --list-engines)")

    input_file = Path(args.input)

    if not input_file.exists():
        print(f"❌ Error: Input file not found: {input_file}")
        sys.exit(1)

    if not args.pdf and not args.docx:
        parser.error("Specify --pdf and/or --docx output file")

    # Generate outputs
    success = True

    if args.pdf:
        pdf_success = export_pdf(input_file, Path(args.pdf), engine=args.engine)
        success = success and pdf_success

    if args.docx:
        docx_success = export_docx(input_file, Path(args.docx))
        success = success and docx_success

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
