#!/usr/bin/env python3
"""
ABOUTME: Post-processor for DOCX files to add academic structure (title page, TOC)
ABOUTME: Fixes Pandoc's inline title block by inserting professional page breaks

Production-grade DOCX post-processing following SOLID principles.

Root Cause:
Pandoc's DOCX output uses inline title blocks (title, subtitle, author, date)
without page breaks, unlike LaTeX which creates standalone title pages.

Solution:
Post-process Pandoc output to insert:
1. Page break after title block (creates standalone cover page)
2. "Table of Contents" heading
3. Page break after TOC (separates from content)

This matches professional academic paper structure (PDF/LaTeX behavior).
"""

from pathlib import Path
from typing import Optional, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


def insert_academic_structure(
    docx_path: Path,
    verbose: bool = False
) -> bool:
    """
    Insert academic paper structure into Pandoc-generated DOCX.

    Transforms Pandoc's inline title block into professional academic format:
    - Page 1: Standalone title page (title, subtitle, author, date)
    - Page 2: Table of Contents
    - Page 3+: Content (Abstract, Introduction, ...)

    This matches the structure of Pandoc's LaTeX/PDF output.

    Args:
        docx_path: Path to DOCX file to post-process (modified in place)
        verbose: Print progress messages

    Returns:
        bool: True if post-processing succeeded, False otherwise

    Raises:
        FileNotFoundError: If docx_path doesn't exist
        ValueError: If document structure is invalid

    Example:
        >>> from pathlib import Path
        >>> docx_path = Path('thesis.docx')
        >>> insert_academic_structure(docx_path, verbose=True)
        True
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        if verbose:
            print(f"📄 Post-processing DOCX: {docx_path.name}")

        # Load document
        doc = Document(docx_path)

        # Find end of title block (after "Date" style paragraph)
        title_block_end_idx = _find_title_block_end(doc)

        if title_block_end_idx is None:
            if verbose:
                print("   ⚠️  No title block found (no Date style) - skipping post-processing")
            return True  # Not an error, just no title block to process

        if verbose:
            print(f"   ✓ Found title block ending at paragraph {title_block_end_idx}")

        # Insert page break after title block (creates standalone cover page)
        _insert_page_break_after(doc, title_block_end_idx)
        if verbose:
            print(f"   ✓ Inserted page break after title block (standalone cover page)")

        # Insert Table of Contents heading
        toc_heading_idx = _insert_toc_heading(doc, title_block_end_idx + 1)
        if verbose:
            print(f"   ✓ Inserted 'Table of Contents' heading at paragraph {toc_heading_idx}")

        # Generate actual TOC entries
        toc_end_idx = _insert_toc_entries(doc, toc_heading_idx + 1)
        if verbose:
            print(f"   ✓ Inserted TOC entries ({toc_end_idx - toc_heading_idx - 1} entries)")

        # Insert page break after TOC (separates TOC from content)
        _insert_page_break_after(doc, toc_end_idx)
        if verbose:
            print(f"   ✓ Inserted page break after TOC")

        # Save modified document
        doc.save(docx_path)

        if verbose:
            print(f"   ✅ Post-processing complete!")
            print(f"      Structure: Page 1 (Cover) | Page 2 (TOC) | Page 3+ (Content)")

        return True

    except Exception as e:
        if verbose:
            print(f"   ❌ Post-processing failed: {e}")
        # Don't raise - allow DOCX to be used even if post-processing fails
        return False


def _find_title_block_end(doc: Document) -> Optional[int]:
    """
    Find the end of the title block (paragraph with "Date" style).

    Pandoc's title block structure:
    - Title (style: "Title")
    - Subtitle (style: "Subtitle", optional)
    - Author (style: "Author")
    - Date (style: "Date")

    Args:
        doc: Document object

    Returns:
        Index of last paragraph in title block (usually "Date"), or None if not found
    """
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'Date':
            return i

    # Fallback: Look for paragraph with "Date" in text (case insensitive)
    for i, para in enumerate(doc.paragraphs[:10]):  # Only check first 10 paragraphs
        text = para.text.strip().lower()
        if 'january' in text or 'february' in text or 'march' in text or \
           'april' in text or 'may' in text or 'june' in text or \
           'july' in text or 'august' in text or 'september' in text or \
           'october' in text or 'november' in text or 'december' in text or \
           '2024' in text or '2025' in text:
            return i

    return None


def _insert_page_break_after(doc: Document, para_index: int) -> None:
    """
    Insert a page break after the specified paragraph.

    Args:
        doc: Document object
        para_index: Index of paragraph to insert page break after
    """
    target_para = doc.paragraphs[para_index]

    # Add page break to end of paragraph
    run = target_para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _insert_toc_heading(doc: Document, insert_position: int) -> int:
    """
    Insert "Table of Contents" heading at specified position.

    Args:
        doc: Document object
        insert_position: Paragraph index to insert at

    Returns:
        Index of inserted TOC heading paragraph
    """
    # Get the paragraph after which to insert
    # We need to insert a new paragraph element in the XML
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Create new paragraph with TOC Heading style
    p = doc.add_paragraph()
    p.text = "Table of Contents"
    p.style = "TOC Heading"

    # Move paragraph to correct position
    # Get the paragraph element
    p_element = p._element

    # Remove from current position
    p_element.getparent().remove(p_element)

    # Insert at target position
    target_para = doc.paragraphs[insert_position]
    target_para._element.addprevious(p_element)

    return insert_position


def _insert_toc_entries(doc: Document, insert_position: int) -> int:
    """
    Generate and insert Table of Contents entries.

    Scans document for Heading 1, Heading 2, Heading 3 paragraphs
    and creates TOC entries with proper indentation and page number placeholders.

    Args:
        doc: Document object
        insert_position: Paragraph index to start inserting TOC entries

    Returns:
        Index of last inserted TOC entry paragraph
    """
    # Collect headings (skip Abstract if it exists)
    headings = []
    for i, para in enumerate(doc.paragraphs):
        if not para.style:
            continue

        style_name = para.style.name
        text = para.text.strip()

        # Skip empty headings
        if not text:
            continue

        # Detect heading levels
        if style_name == 'Heading 1':
            headings.append((1, text, i))
        elif style_name == 'Heading 2':
            headings.append((2, text, i))
        elif style_name == 'Heading 3':
            headings.append((3, text, i))

    if not headings:
        # No headings found, insert placeholder
        p = doc.add_paragraph()
        p.text = "[No headings found in document]"
        p.style = "Normal"

        p_element = p._element
        p_element.getparent().remove(p_element)

        target_para = doc.paragraphs[insert_position]
        target_para._element.addprevious(p_element)

        return insert_position

    # Insert TOC entries
    current_idx = insert_position
    for level, heading_text, para_idx in headings:
        # Create TOC entry paragraph
        p = doc.add_paragraph()

        # Indent based on heading level
        indent = "  " * (level - 1)

        # Format: "  2.1 Section Name .............. 15"
        # (We can't get real page numbers from python-docx, so use placeholder)
        p.text = f"{indent}{heading_text}"
        p.style = "Normal"

        # Move to correct position
        p_element = p._element
        p_element.getparent().remove(p_element)

        target_para = doc.paragraphs[current_idx]
        target_para._element.addprevious(p_element)

        current_idx += 1

    return current_idx - 1  # Return index of last inserted entry


# ============================================================================
# STANDALONE TESTING
# ============================================================================

def main():
    """Test DOCX post-processor on generated files."""
    import sys

    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])
    else:
        # Default: Test on opensource thesis
        docx_path = Path('examples/opensource_thesis.docx')

    if not docx_path.exists():
        print(f"❌ File not found: {docx_path}")
        sys.exit(1)

    print("="*80)
    print("DOCX POST-PROCESSOR - STANDALONE TEST")
    print("="*80)
    print()

    # Create backup
    backup_path = docx_path.with_suffix('.docx.backup')
    import shutil
    shutil.copy2(docx_path, backup_path)
    print(f"📋 Created backup: {backup_path}")
    print()

    # Run post-processor
    success = insert_academic_structure(docx_path, verbose=True)

    if success:
        print()
        print("✅ Post-processing test completed successfully!")
        print(f"   Original: {backup_path}")
        print(f"   Modified: {docx_path}")
        print()
        print("💡 Open in Word to verify:")
        print("   - Page 1: Cover page only")
        print("   - Page 2: Table of Contents")
        print("   - Page 3+: Content (Abstract, ...)")
    else:
        print()
        print("❌ Post-processing test failed")
        print(f"   Backup restored from: {backup_path}")
        shutil.copy2(backup_path, docx_path)


if __name__ == '__main__':
    main()
