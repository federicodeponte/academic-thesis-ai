#!/usr/bin/env python3
"""
ABOUTME: Legacy export utilities for markdown to PDF, Word, and LaTeX conversion
ABOUTME: For production use, prefer utils/export_professional.py with multi-engine support
"""

import argparse
import os
import sys
import markdown
from pathlib import Path

# Use centralized logging system
from utils.logging_config import get_logger
logger = get_logger(__name__)

# Verification infrastructure
try:
    from citation_verifier import CitationVerifier
    from fact_checker import FactChecker
    VERIFICATION_AVAILABLE = True
except ImportError:
    # Try relative import if running from utils/
    try:
        import sys
        sys.path.insert(0, os.path.dirname(__file__))
        from citation_verifier import CitationVerifier
        from fact_checker import FactChecker
        VERIFICATION_AVAILABLE = True
    except ImportError:
        VERIFICATION_AVAILABLE = False
        logger.warning("Verification infrastructure not available. Use --skip-verification to bypass.")

try:
    from weasyprint import HTML, CSS
    WEASY_AVAILABLE = True
except ImportError:
    WEASY_AVAILABLE = False
    logger.warning("WeasyPrint not installed. PDF export unavailable.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word export unavailable.")


def verify_before_export(input_file, force=False, skip_verification=False):
    """
    Verify academic integrity before export (Quality Gate)

    Returns:
        True if verification passed or was skipped/forced
        False if verification failed and not forced
    """
    if skip_verification:
        logger.warning("="*70)
        logger.warning("Skipping verification (--skip-verification)")
        logger.warning("Export may contain unverified citations and claims.")
        logger.warning("="*70)
        return True

    if not VERIFICATION_AVAILABLE:
        logger.warning("="*70)
        logger.warning("Verification infrastructure not available")
        logger.warning("Install dependencies or use --skip-verification")
        logger.warning("="*70)
        return False

    logger.info("="*70)
    logger.info("ACADEMIC INTEGRITY VERIFICATION")
    logger.info("="*70)

    # Citation verification
    logger.info("Running citation verification...")
    citation_verifier = CitationVerifier(use_cache=True)
    citation_report = citation_verifier.verify_file(input_file)
    citation_report.print_summary()

    # Fact-checking
    logger.info("Running fact-check...")
    fact_checker = FactChecker()
    fact_report = fact_checker.check_file(input_file)
    fact_report.print_summary()

    # Overall assessment
    citation_passed = citation_report.passes_threshold(95.0)
    fact_passed = fact_report.passes_threshold(95.0)

    logger.info("="*70)
    logger.info("VERIFICATION SUMMARY")
    logger.info("="*70)
    logger.info(f"Citations: {'PASS' if citation_passed else 'FAIL'} ({citation_report.verification_rate:.1f}%)")
    logger.info(f"Fact-check: {'PASS' if fact_passed else 'FAIL'} ({fact_report.citation_rate:.1f}%)")

    if citation_passed and fact_passed:
        logger.info("VERIFICATION PASSED - Ready for export")
        logger.info("="*70)
        return True
    else:
        logger.error("VERIFICATION FAILED - Academic integrity issues detected")
        logger.error("="*70)

        if force:
            logger.warning("Proceeding with export anyway (--force)")
            logger.warning("This document has NOT met academic standards.")
            return True
        else:
            logger.error("Export BLOCKED. Fix issues or use --force to override.")
            logger.error("Use --skip-verification to bypass checks entirely.")
            return False


def markdown_to_html(md_content):
    """Convert Markdown to HTML"""
    md = markdown.Markdown(extensions=[
        'extra',
        'codehilite',
        'tables',
        'toc',
        'footnotes'
    ])
    return md.convert(md_content)


def export_pdf(input_file, output_file):
    """Export Markdown to PDF using WeasyPrint"""
    if not WEASY_AVAILABLE:
        logger.error("WeasyPrint not installed. Run: pip install weasyprint")
        return False
    
    # Read Markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert to HTML
    html_content = markdown_to_html(md_content)
    
    # Academic CSS styling
    css = CSS(string='''
        @page {
            size: A4;
            margin: 2.5cm;
        }
        body {
            font-family: "Times New Roman", serif;
            font-size: 12pt;
            line-height: 1.6;
            text-align: justify;
        }
        h1 {
            font-size: 16pt;
            font-weight: bold;
            margin-top: 24pt;
            margin-bottom: 12pt;
        }
        h2 {
            font-size: 14pt;
            font-weight: bold;
            margin-top: 18pt;
            margin-bottom: 9pt;
        }
        h3 {
            font-size: 12pt;
            font-weight: bold;
            margin-top: 12pt;
            margin-bottom: 6pt;
        }
        p {
            margin-bottom: 12pt;
            text-indent: 0.5in;
        }
        blockquote {
            margin-left: 0.5in;
            margin-right: 0.5in;
            font-style: italic;
        }
        code {
            font-family: "Courier New", monospace;
            background-color: #f5f5f5;
            padding: 2px 4px;
        }
        pre {
            font-family: "Courier New", monospace;
            background-color: #f5f5f5;
            padding: 10px;
            overflow-x: auto;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 12pt 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f5f5f5;
            font-weight: bold;
        }
    ''')
    
    # Generate PDF
    HTML(string=html_content).write_pdf(output_file, stylesheets=[css])
    logger.info(f"PDF exported successfully: {output_file}")
    return True


def export_docx(input_file, output_file):
    """Export Markdown to Word (.docx)"""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return False
    
    # Read Markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Parse Markdown (simple line-by-line conversion)
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. '):
            doc.add_paragraph(line[3:], style='List Number')
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
        
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
        
        i += 1
    
    # Save
    doc.save(output_file)
    logger.info(f"Word document exported successfully: {output_file}")
    return True


def export_latex(input_file, output_file):
    """Export Markdown to LaTeX"""
    # Read Markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # LaTeX preamble
    latex_content = r'''\documentclass[12pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{times}
\usepackage[margin=1in]{geometry}
\usepackage{graphicx}
\usepackage{hyperref}
\usepackage{cite}
\usepackage{amsmath}
\usepackage{booktabs}

\title{Your Paper Title}
\author{Your Name}
\date{\today}

\begin{document}

\maketitle

'''
    
    # Simple Markdown to LaTeX conversion
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        
        if not line:
            latex_content += '\n'
            continue
        
        # Headings
        if line.startswith('# '):
            latex_content += f'\\section{{{line[2:]}}}\n'
        elif line.startswith('## '):
            latex_content += f'\\subsection{{{line[3:]}}}\n'
        elif line.startswith('### '):
            latex_content += f'\\subsubsection{{{line[4:]}}}\n'
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            if '\\begin{itemize}' not in latex_content.split('\n')[-2]:
                latex_content += '\\begin{itemize}\n'
            latex_content += f'\\item {line[2:]}\n'
        
        # Regular text
        else:
            # Close itemize if needed
            if '\\begin{itemize}' in latex_content and '\\end{itemize}' not in latex_content.split('\n')[-2]:
                latex_content += '\\end{itemize}\n'
            
            # Escape special characters
            line = line.replace('&', '\\&').replace('%', '\\%').replace('$', '\\$')
            latex_content += line + '\n\n'
    
    latex_content += '\n\\end{document}\n'
    
    # Save
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(latex_content)

    logger.info(f"LaTeX exported successfully: {output_file}")
    logger.info(f"Compile with: pdflatex {output_file}")
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Export Markdown to PDF, Word, or LaTeX with academic integrity verification',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Normal export (with verification)
  python export.py --format pdf --output paper.pdf final_thesis.md

  # Skip verification (not recommended)
  python export.py --format pdf --output paper.pdf --skip-verification final_thesis.md

  # Force export despite verification failures
  python export.py --format pdf --output paper.pdf --force final_thesis.md

Academic Integrity:
  By default, this tool verifies:
  - Citations exist and are accessible (CrossRef, arXiv)
  - Quantitative claims are properly cited
  - No hallucinated references (95% threshold)

  Use --force or --skip-verification at your own risk.
        '''
    )
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('--format', '-f', choices=['pdf', 'docx', 'latex'], required=True,
                        help='Output format')
    parser.add_argument('--output', '-o', required=True,
                        help='Output file path')
    parser.add_argument('--force', action='store_true',
                        help='Force export even if verification fails (not recommended)')
    parser.add_argument('--skip-verification', action='store_true',
                        help='Skip verification entirely (not recommended)')

    args = parser.parse_args()

    # Check input file exists
    if not os.path.exists(args.input):
        logger.error(f"Input file not found: {args.input}")
        sys.exit(1)

    # Verify academic integrity before export (Quality Gate)
    verification_passed = verify_before_export(
        args.input,
        force=args.force,
        skip_verification=args.skip_verification
    )

    if not verification_passed:
        logger.error("Export BLOCKED due to verification failure.")
        logger.error("Fix issues, use --force, or use --skip-verification")
        sys.exit(1)

    # Create output directory if needed
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # Export
    success = False
    if args.format == 'pdf':
        success = export_pdf(args.input, args.output)
    elif args.format == 'docx':
        success = export_docx(args.input, args.output)
    elif args.format == 'latex':
        success = export_latex(args.input, args.output)

    if not success:
        sys.exit(1)


if __name__ == '__main__':
    main()
